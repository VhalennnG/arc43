import os
import re
from typing import List, Dict, Any
from docx import Document
from src.analyzers.base import AbstractAnalyzer
from src.fields.models import FormField, FieldType

# Named constant for cell-based checkbox width threshold
# 300,000 EMU ≈ 0.33 inches — cells narrower than this are treated as checkbox tick-boxes
CELL_CHECKBOX_WIDTH_THRESHOLD_EMU = 300000

class DocxAnalyzer(AbstractAnalyzer):
    """
    Layout analyzer for Word (DOCX) documents.
    Detects table cells, paragraphs, and checkboxes (symbol and SDT content controls).
    """

    def clean_label(self, text: str) -> str:
        # Remove colons, underscores, dots, hyphens, and whitespace
        text = re.sub(r'[:：_．\.\-\s]+$', '', text)
        return text.strip()

    def is_empty_field_text(self, text: str) -> bool:
        # Matches if string is empty or only contains spaces, underscores, dots, or hyphens
        return re.match(r'^[_\.\-\s\u2026]*$', text) is not None

    def detect_sdt_checkboxes(self, doc: Document) -> List[FormField]:
        fields = []
        # Find all structured document tags in the document XML using local-name() to be namespace-agnostic
        sdt_elements = doc.element.xpath('.//*[local-name()="sdt"]')
        
        for idx, sdt in enumerate(sdt_elements):
            # Check if this SDT contains a checkbox definition
            checkbox_nodes = sdt.xpath('.//*[local-name()="checkbox"]')
            if checkbox_nodes:
                # It's an SDT checkbox
                # Try to extract the title/alias as the label
                alias_nodes = sdt.xpath('.//*[local-name()="sdtPr"]/*[local-name()="alias"]/@*[local-name()="val"]')
                tag_nodes = sdt.xpath('.//*[local-name()="sdtPr"]/*[local-name()="tag"]/@*[local-name()="val"]')
                
                label = ""
                if alias_nodes:
                    label = str(alias_nodes[0])
                elif tag_nodes:
                    label = str(tag_nodes[0])
                else:
                    label = f"Checkbox_SDT_{idx}"
                    
                field_id = f"sdt_checkbox_{idx}"
                # Store only the serializable index, not the lxml Element (#15)
                field = FormField(
                    id=field_id,
                    label=label,
                    field_type=FieldType.CHECKBOX,
                    checkbox_kind="sdt",
                    metadata={"requires_field_update": True, "sdt_index": idx}
                )
                fields.append(field)
        return fields

    def detect_symbol_checkboxes_for_paragraph(self, p, p_idx: int, context_labels: List[str],
                                                table_idx: int = None, row_idx: int = None, col_idx: int = None) -> List[FormField]:
        """
        Detects symbol checkboxes (☐/□) within a single paragraph.
        Accepts the actual paragraph index to avoid ID collisions (#6).
        """
        fields = []
        p_text = p.text
        first_box_idx = p_text.find("☐")
        if first_box_idx == -1:
            first_box_idx = p_text.find("□")
            
        if first_box_idx == -1:
            return fields
            
        # Extract header label before the first checkbox
        group_header = self.clean_label(p_text[:first_box_idx])
        
        for run_idx, run in enumerate(p.runs):
            box_char = None
            if "☐" in run.text:
                box_char = "☐"
            elif "□" in run.text:
                box_char = "□"
                
            if box_char:
                # Skip if run is inside an SDT control to avoid double-detection (Phase 6)
                if run.element.xpath('ancestor::*[local-name()="sdt"]'):
                    continue
                    
                # Find option label after the box
                text_after = run.text.split(box_char, 1)[1].strip()
                if not text_after:
                    # Look in subsequent runs
                    for next_run in p.runs[run_idx+1:]:
                        next_text = next_run.text.strip()
                        if next_text:
                            if "☐" in next_text or "□" in next_text:
                                break
                            text_after = next_text
                            break
                            
                option_label = self.clean_label(text_after)
                if not option_label:
                    option_label = f"Opsi_{run_idx}"
                    
                context = list(context_labels)
                if group_header:
                    context.append(group_header)
                    
                if table_idx is not None:
                    field_id = f"table_{table_idx}_row_{row_idx}_col_{col_idx}_p_{p_idx}_run_{run_idx}"
                else:
                    field_id = f"paragraph_{p_idx}_run_{run_idx}"
                    
                field = FormField(
                    id=field_id,
                    label=option_label,
                    field_type=FieldType.CHECKBOX,
                    context_labels=context,
                    table_index=table_idx,
                    row=row_idx,
                    column=col_idx,
                    paragraph_index=p_idx if table_idx is None else None,
                    checkbox_kind="symbol"
                )
                fields.append(field)
        return fields

    def detect_symbol_checkboxes(self, paragraphs, context_labels: List[str], 
                                 table_idx: int = None, row_idx: int = None, col_idx: int = None) -> List[FormField]:
        """Legacy wrapper: iterates paragraphs in a table cell context."""
        fields = []
        for p_idx, p in enumerate(paragraphs):
            fields.extend(self.detect_symbol_checkboxes_for_paragraph(
                p, p_idx, context_labels, table_idx, row_idx, col_idx
            ))
        return fields

    def analyze_table_row(self, t_idx: int, r_idx: int, row) -> List[FormField]:
        fields = []
        num_cols = len(row.cells)
        visited = set()
        
        # First pass: find all cell-based checkboxes (narrow empty cell followed by option text)
        for c_idx in range(num_cols - 1):
            if c_idx in visited:
                continue
                
            cell = row.cells[c_idx]
            next_cell = row.cells[c_idx + 1]
            
            # Check if current cell is empty and narrow, and next cell contains choice text
            if self.is_empty_field_text(cell.text) and cell.width < CELL_CHECKBOX_WIDTH_THRESHOLD_EMU:
                next_text = next_cell.text.strip()
                if next_text and not self.is_empty_field_text(next_text):
                    # It's a cell-based checkbox!
                    option_label = self.clean_label(next_text)
                    
                    # Find the group header by scanning backward
                    group_header = ""
                    for prev_idx in range(c_idx - 1, -1, -1):
                        prev_text = row.cells[prev_idx].text.strip()
                        if prev_text and prev_text not in (":", "：") and not self.is_empty_field_text(prev_text):
                            group_header = self.clean_label(prev_text)
                            break
                            
                    fields.append(FormField(
                        id=f"table_{t_idx}_row_{r_idx}_col_{c_idx}",
                        label=option_label,
                        field_type=FieldType.CHECKBOX,
                        checkbox_kind="cell",
                        context_labels=[group_header] if group_header else [],
                        table_index=t_idx,
                        row=r_idx,
                        column=c_idx
                    ))
                    visited.add(c_idx)
                    visited.add(c_idx + 1)
                    
        # Second pass: find text input fields
        for c_idx in range(num_cols):
            if c_idx in visited:
                continue
                
            cell = row.cells[c_idx]
            cell_text = cell.text.strip()
            
            if not cell_text or cell_text in (":", "："):
                continue
                
            label = self.clean_label(cell_text)
            
            # Check if this cell is followed by an empty cell
            input_idx = c_idx + 1
            if input_idx < num_cols and row.cells[input_idx].text.strip() in (":", "："):
                input_idx += 1
                
            if input_idx < num_cols:
                input_cell = row.cells[input_idx]
                if input_idx not in visited and self.is_empty_field_text(input_cell.text):
                    # It's an empty cell. Since it wasn't registered as a checkbox, it's a text input!
                    fields.append(FormField(
                        id=f"table_{t_idx}_row_{r_idx}_col_{input_idx}",
                        label=label,
                        field_type=FieldType.TEXT,
                        table_index=t_idx,
                        row=r_idx,
                        column=input_idx
                    ))
                    visited.add(input_idx)
                    # Skip subsequent consecutive empty cells in the same row
                    skip_idx = input_idx + 1
                    while skip_idx < num_cols and self.is_empty_field_text(row.cells[skip_idx].text):
                        visited.add(skip_idx)
                        skip_idx += 1
                        
        return fields

    def analyze(self, file_path: str) -> List[FormField]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX template not found: {file_path}")
            
        doc = Document(file_path)
        fields = []
        
        # 1. Detect SDT checkboxes first
        fields.extend(self.detect_sdt_checkboxes(doc))
        
        # 2. Detect paragraph-based fields
        for p_idx, p in enumerate(doc.paragraphs):
            p_text = p.text.strip()
            if not p_text:
                continue
                
            # If the paragraph contains SDT controls, skip it for plain paragraph fields (handled in step 1)
            if p._p.xpath('.//*[local-name()="sdt"]'):
                continue
                
            # If it has checkbox symbol, parse as symbol checkboxes — pass actual p_idx (#6)
            if "☐" in p_text or "□" in p_text:
                fields.extend(self.detect_symbol_checkboxes_for_paragraph(p, p_idx, []))
                continue
                
            # Check for label-underline patterns: "Name: _____" or "Name: " at the end
            if ":" in p_text or "：" in p_text:
                delim = ":" if ":" in p_text else "："
                parts = p_text.split(delim, 1)
                label = self.clean_label(parts[0])
                val = parts[1].strip()
                if self.is_empty_field_text(val):
                    fields.append(FormField(
                        id=f"paragraph_{p_idx}",
                        label=label,
                        field_type=FieldType.TEXT,
                        paragraph_index=p_idx
                    ))
            else:
                # Check for "Name ________" without colons
                m = re.search(r'^([^_\.\u2026]+?)\s*([_\.\u2026]{2,})\s*$', p_text)
                if m:
                    label = self.clean_label(m.group(1))
                    fields.append(FormField(
                        id=f"paragraph_{p_idx}",
                        label=label,
                        field_type=FieldType.TEXT,
                        paragraph_index=p_idx
                    ))
                    
        # 3. Detect table-based fields
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                fields.extend(self.analyze_table_row(t_idx, r_idx, row))
                            
        return fields
