import os
import re
from typing import List
from docx import Document
from src.writers.base import DocumentWriter
from src.fields.models import FormField, FieldType

def replace_text_preserving_format(paragraph, new_text: str) -> None:
    """
    Replaces the text content of a paragraph while preserving the formatting
    (bold, italic, font, size) of the first run. Clears subsequent runs.
    This avoids the destructive behavior of `paragraph.text = ...` which
    strips all run-level formatting (#10).
    """
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""

def replace_cell_text_preserving_format(cell, new_text: str) -> None:
    """
    Replaces the text of the first paragraph in a cell while preserving formatting.
    Unlike `cell.text = ...`, this does not delete extra paragraphs or strip fonts.
    """
    if cell.paragraphs:
        replace_text_preserving_format(cell.paragraphs[0], new_text)
    else:
        cell.text = new_text

class DocxWriter(DocumentWriter):
    """
    Word (DOCX) document writer.
    Fills text inputs and checks symbol checkboxes.
    Preserves run-level formatting (bold, italic, font) when writing text.
    """

    def is_checked(self, field: FormField) -> bool:
        """
        Determines if a checkbox should be checked based on the answer.
        """
        if not field.answer:
            return False
            
        ans = field.answer
        if isinstance(ans, list):
            # Check if option label is in the answer list
            return any(
                str(a).strip().lower() in field.label.strip().lower() or 
                field.label.strip().lower() in str(a).strip().lower() 
                for a in ans
            )
            
        ans_str = str(ans).strip().lower()
        label_str = field.label.strip().lower()
        
        # If answer matches option label, or is a substring, or is explicitly truthy
        return (
            ans_str == label_str or 
            ans_str in label_str or 
            label_str in ans_str or 
            ans_str in ("yes", "true", "1", "ya", "checked")
        )

    def _extract_group_header(self, paragraph) -> str:
        """Extracts the group header text before the first checkbox symbol in a paragraph."""
        p_text = paragraph.text
        idx = p_text.find("☐")
        if idx == -1:
            idx = p_text.find("□")
        if idx > 0:
            header = p_text[:idx].strip().rstrip(":：").strip()
            return header.lower()
        return ""

    def _context_matches(self, field: FormField, paragraph) -> bool:
        """
        Verifies that the context_labels of a field match the group header
        of the candidate paragraph, for disambiguation (#7).
        """
        if not field.context_labels:
            return True  # No context constraint — accept any match
        candidate_header = self._extract_group_header(paragraph)
        if not candidate_header:
            return True  # No header to compare — accept
        return any(cl.lower() in candidate_header or candidate_header in cl.lower() for cl in field.context_labels)

    def fill_paragraphs(self, paragraphs, field: FormField) -> bool:
        """
        Attempts to write text value or toggle checkbox in paragraphs.
        Returns True if the field was handled.
        """
        # Look for target paragraph by index if set
        if field.paragraph_index is not None and field.paragraph_index < len(paragraphs):
            p = paragraphs[field.paragraph_index]
            
            if field.field_type == FieldType.CHECKBOX and field.checkbox_kind == "symbol":
                # Find the run and toggle it
                for run in p.runs:
                    if "☐" in run.text or "□" in run.text:
                        # Only toggle if option matches the target run text
                        box_char = "☐" if "☐" in run.text else "□"
                        run_text_after = run.text.split(box_char, 1)[1].strip()
                        if field.label.lower() in run.text.lower() or (not run_text_after and any(field.label.lower() in r.text.lower() for r in p.runs)):
                            if self.is_checked(field):
                                run.text = run.text.replace(box_char, "☑")
                            return True
            else:
                # Text field: replace underscores or append text
                # Find the run containing underscores/dots, or append to the end
                replaced = False
                for run in p.runs:
                    if re.search(r'[_\.\u2026]{2,}', run.text):
                        run.text = re.sub(r'[_\.\u2026]{2,}', str(field.answer or ""), run.text, count=1)
                        replaced = True
                        break
                if not replaced:
                    # Append answer preserving format of existing runs
                    if p.runs:
                        p.runs[-1].text = p.runs[-1].text + " " + str(field.answer or "")
                    else:
                        p.add_run(" " + str(field.answer or ""))
                return True
                
        # If index is not set or not matched, fall back to searching label/context
        if field.field_type == FieldType.CHECKBOX and field.checkbox_kind == "symbol":
            for p in paragraphs:
                if "☐" in p.text or "□" in p.text:
                    # Verify context_labels match the group header of this paragraph (#7)
                    if not self._context_matches(field, p):
                        continue
                    for run_idx, run in enumerate(p.runs):
                        box_char = "☐" if "☐" in run.text else "□"
                        if box_char in run.text:
                            # Match label
                            text_after = run.text.split(box_char, 1)[1].strip()
                            if not text_after:
                                for next_run in p.runs[run_idx+1:]:
                                    if next_run.text.strip():
                                        text_after = next_run.text.strip()
                                        break
                            if field.label.lower() in text_after.lower():
                                if self.is_checked(field):
                                    run.text = run.text.replace(box_char, "☑")
                                return True
        return False

    def fill(self, source_path: str, fields: List[FormField], output_path: str) -> None:
        if not os.path.exists(source_path):
            raise FileNotFoundError(f"DOCX template not found: {source_path}")
            
        doc = Document(source_path)
        
        for field in fields:
            # Skip if no answer is provided
            if field.answer is None:
                continue
                
            # Case A: SDT checkbox content control
            if field.field_type == FieldType.CHECKBOX and field.checkbox_kind == "sdt":
                sdt_elements = doc.element.xpath('.//*[local-name()="sdt"]')
                m = re.match(r'sdt_checkbox_(\d+)', field.id)
                if m:
                    idx = int(m.group(1))
                    if idx < len(sdt_elements):
                        sdt = sdt_elements[idx]
                        checked = self.is_checked(field)
                        
                        # 1. Update <w14:checked w14:val="1/0"/>
                        checked_nodes = sdt.xpath('.//*[local-name()="checked"]')
                        if checked_nodes:
                            checked_node = checked_nodes[0]
                            val_attr = None
                            for key in checked_node.attrib.keys():
                                if key.endswith("val"):
                                    val_attr = key
                                    break
                            if val_attr:
                                checked_node.set(val_attr, "1" if checked else "0")
                                
                        # 2. Determine display character and font from checkedState/uncheckedState (#14)
                        check_char = "☒"
                        uncheck_char = "☐"
                        check_font = None
                        uncheck_font = None
                        
                        checked_state_nodes = sdt.xpath('.//*[local-name()="checkedState"]')
                        unchecked_state_nodes = sdt.xpath('.//*[local-name()="uncheckedState"]')
                        
                        if checked_state_nodes:
                            cs = checked_state_nodes[0]
                            val = None
                            font = None
                            for key, value in cs.attrib.items():
                                if key.endswith("val"):
                                    val = value
                                elif key.endswith("font"):
                                    font = value
                            if val:
                                try:
                                    check_char = chr(int(val, 16))
                                except Exception:
                                    pass
                            check_font = font
                            
                        if unchecked_state_nodes:
                            us = unchecked_state_nodes[0]
                            val = None
                            font = None
                            for key, value in us.attrib.items():
                                if key.endswith("val"):
                                    val = value
                                elif key.endswith("font"):
                                    font = value
                            if val:
                                try:
                                    uncheck_char = chr(int(val, 16))
                                except Exception:
                                    pass
                            uncheck_font = font
                                
                        # 3. Update display character in <w:sdtContent>
                        t_nodes = sdt.xpath('.//*[local-name()="sdtContent"]//*[local-name()="t"]')
                        if t_nodes:
                            t_nodes[0].text = check_char if checked else uncheck_char
                            
                        # 4. Apply the corresponding font to the run's rPr/rFonts (#14)
                        target_font = check_font if checked else uncheck_font
                        if target_font:
                            r_nodes = sdt.xpath('.//*[local-name()="sdtContent"]//*[local-name()="r"]')
                            if r_nodes:
                                r_node = r_nodes[0]
                                # Find or create rPr
                                rpr_nodes = r_node.xpath('./*[local-name()="rPr"]')
                                if rpr_nodes:
                                    rpr = rpr_nodes[0]
                                else:
                                    from lxml import etree
                                    nsmap = r_node.nsmap
                                    w_ns = nsmap.get('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
                                    rpr = etree.SubElement(r_node, f'{{{w_ns}}}rPr')
                                    # Insert rPr as first child
                                    r_node.insert(0, rpr)
                                    
                                # Find or create rFonts
                                rfonts_nodes = rpr.xpath('./*[local-name()="rFonts"]')
                                if rfonts_nodes:
                                    rfonts = rfonts_nodes[0]
                                else:
                                    from lxml import etree
                                    nsmap = rpr.nsmap
                                    w_ns = nsmap.get('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
                                    rfonts = etree.SubElement(rpr, f'{{{w_ns}}}rFonts')
                                    
                                # Set all font attributes to the target font
                                for key in list(rfonts.attrib.keys()):
                                    if any(key.endswith(attr) for attr in ("ascii", "hAnsi", "eastAsia", "cs")):
                                        rfonts.set(key, target_font)

            # Case B: Field is in a table
            elif field.table_index is not None and field.table_index < len(doc.tables):
                table = doc.tables[field.table_index]
                if field.row is not None and field.row < len(table.rows):
                    row = table.rows[field.row]
                    if field.column is not None and field.column < len(row.cells):
                        cell = row.cells[field.column]
                        
                        if field.field_type == FieldType.CHECKBOX and field.checkbox_kind == "symbol":
                            # Checkbox in cell paragraphs
                            self.fill_paragraphs(cell.paragraphs, field)
                        elif field.field_type == FieldType.CHECKBOX and field.checkbox_kind == "cell":
                            # Checkbox in a table cell (narrow empty cell)
                            if self.is_checked(field):
                                replace_cell_text_preserving_format(cell, "X")
                        else:
                            # Text field in cell: replace text preserving format (#10)
                            replace_cell_text_preserving_format(cell, str(field.answer))
                                
            # Case C: Field is in body paragraphs
            elif field.paragraph_index is not None:
                self.fill_paragraphs(doc.paragraphs, field)
                
        # Save output to destination
        doc.save(output_path)
