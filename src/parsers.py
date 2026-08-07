import os
from typing import Dict, Any
from pypdf import PdfReader
from docx import Document
from openpyxl import load_workbook

def parse_pdf(file_path: str) -> tuple[str, str]:
    """
    Parses a text-based PDF file and extracts text page-by-page.
    Falls back to native macOS Vision OCR if programmatic extraction returns empty.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"PDF file not found: {file_path}")
        
    reader = PdfReader(file_path)
    text_parts = []
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text and page_text.strip():
            text_parts.append(page_text.strip())
            
    extracted_text = "\n".join(text_parts).strip()
    
    # Check if empty -> Trigger native OCR fallback
    if not extracted_text:
        print(f"Programmatic PDF extraction returned empty for {os.path.basename(file_path)}. Triggering native macOS Vision OCR fallback...")
        from src import ocr
        ocr_text = ocr.recognize_pdf_text_via_ocr(file_path)
        return ocr_text, "ocr"
        
    return extracted_text, "parser"

def parse_docx(file_path: str) -> str:
    """
    Parses a Word document (DOCX) and extracts text from paragraphs and tables.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
    doc = Document(file_path)
    text_parts = []
    
    # Extract from paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            text_parts.append(p.text.strip())
            
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            # De-duplicate adjacent identical cell values caused by merged cells
            clean_row_text = []
            for val in row_text:
                if val and (not clean_row_text or clean_row_text[-1] != val):
                    clean_row_text.append(val)
            if clean_row_text:
                text_parts.append(" | ".join(clean_row_text))
                
    return "\n".join(text_parts)

def parse_xlsx(file_path: str) -> str:
    """
    Parses an Excel spreadsheet (XLSX) and extracts non-empty row values.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"XLSX file not found: {file_path}")
        
    wb = load_workbook(file_path, data_only=True)
    text_parts = []
    
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        text_parts.append(f"--- Sheet: {sheet_name} ---")
        for row in ws.iter_rows(values_only=True):
            # Format row values as string, skipping empty cells
            row_str = " | ".join(str(val).strip() for val in row if val is not None)
            if row_str.strip():
                text_parts.append(row_str)
                
    return "\n".join(text_parts)

def parse_document(file_path: str) -> Dict[str, Any]:
    """
    Detects file extension and routes to appropriate parser.
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        raw_text, method = parse_pdf(file_path)
        return {
            "raw_text": raw_text,
            "source_type": "pdf",
            "extraction_method": method
        }
    elif ext in (".docx", ".doc"):
        return {
            "raw_text": parse_docx(file_path),
            "source_type": "docx",
            "extraction_method": "parser"
        }
    elif ext in (".xlsx", ".xls"):
        return {
            "raw_text": parse_xlsx(file_path),
            "source_type": "xlsx",
            "extraction_method": "parser"
        }
    elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff"):
        return {
            "raw_text": "",
            "source_type": "image",
            "extraction_method": "ocr"
        }
    else:
        raise ValueError(f"Unsupported file format: {ext}")
