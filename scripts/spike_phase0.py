import os
import json
import sys
from docx import Document
from pypdf import PdfReader

# Import src
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from src import llm

def test_docx_checkbox_replacement():
    print("\n--- Test 1: DOCX Checkbox Unicode Replacement ---")
    doc_path = "tests/fixtures/forms/peepl_biodata.docx"
    output_path = "tests/fixtures/forms/peepl_biodata_filled.docx"
    
    if not os.path.exists(doc_path):
        print(f"Error: {doc_path} not found.")
        return False
        
    doc = Document(doc_path)
    
    def replace_in_paragraphs(paragraphs):
        replaced_count = 0
        for p in paragraphs:
            for run in p.runs:
                if "☐" in run.text:
                    # Replace first checkbox for demo
                    run.text = run.text.replace("☐", "☑", 1)
                    replaced_count += 1
        return replaced_count

    # Process document paragraphs
    p_replaced = replace_in_paragraphs(doc.paragraphs)
    
    # Process tables
    t_replaced = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t_replaced += replace_in_paragraphs(cell.paragraphs)
                
    doc.save(output_path)
    print(f"Replaced {p_replaced} in paragraphs, {t_replaced} in tables.")
    
    # Re-open and verify
    doc2 = Document(output_path)
    found_checked = False
    for p in doc2.paragraphs:
        if "☑" in p.text:
            found_checked = True
            print(f"Verified checked symbol in paragraph: '{p.text}'")
            
    for table in doc2.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "☑" in p.text:
                        found_checked = True
                        print(f"Verified checked symbol in table cell: '{p.text}'")
                        
    if os.path.exists(output_path):
        os.remove(output_path)
        
    if found_checked:
        print("Test 1 Result: SUCCESS")
        return True
    else:
        print("Test 1 Result: FAILED")
        return False

def test_pdf_get_fields():
    print("\n--- Test 2: pypdf PdfReader.get_fields() ---")
    acro_path = "tests/fixtures/forms/acroform_sample.pdf"
    static_path = "tests/fixtures/forms/static_sample.pdf"
    
    # Check AcroForm
    reader_acro = PdfReader(acro_path)
    fields_acro = reader_acro.get_fields()
    print(f"AcroForm Fields found: {list(fields_acro.keys()) if fields_acro else None}")
    
    # Check Static
    reader_static = PdfReader(static_path)
    fields_static = reader_static.get_fields()
    print(f"Static Fields found: {list(fields_static.keys()) if fields_static else None}")
    
    if fields_acro and not fields_static:
        print("Test 2 Result: SUCCESS")
        return True
    else:
        print("Test 2 Result: FAILED")
        return False

def test_llm_structured_output():
    print("\n--- Test 3: Local LLM Structured JSON Generation ---")
    prompt = (
        "Output ONLY a valid JSON array of objects representing profile details, matching the format below.\n"
        "Rules:\n"
        "- Do not explain, do not add introductory text, do not use markdown code blocks.\n"
        "- The output must start with '[' and end with ']'.\n\n"
        "Format:\n"
        "[\n"
        '  {"label": "Nama Lengkap", "value": "Vhalentino Gamgenora"}\n'
        "]\n\n"
        "JSON Output:"
    )
    
    print("Sending prompt to local LLM...")
    try:
        raw_output = llm.generate_text(prompt, max_tokens=128, temperature=0.1)
        print("Raw LLM Output:")
        print("---")
        print(raw_output)
        print("---")
        
        # Test parsing
        clean_output = raw_output.strip()
        if clean_output.startswith("```json"):
            clean_output = clean_output[7:]
        elif clean_output.startswith("```"):
            clean_output = clean_output[3:]
        if clean_output.endswith("```"):
            clean_output = clean_output[:-3]
        clean_output = clean_output.strip()
        
        parsed = json.loads(clean_output)
        print(f"Parsed JSON successfully: {parsed}")
        print("Test 3 Result: SUCCESS (Structured generation parsed without grammar)")
        return True
    except Exception as e:
        print(f"Failed to parse JSON: {e}")
        print("Test 3 Result: FAILED / REQUIRES CONSTRAINTS")
        return False

if __name__ == "__main__":
    test_docx_checkbox_replacement()
    test_pdf_get_fields()
    test_llm_structured_output()
