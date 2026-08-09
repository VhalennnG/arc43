import os
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors

def setup_directories():
    os.makedirs("tests/fixtures/forms", exist_ok=True)
    print("Fixture directories verified.")

def generate_docx():
    docx_path = "tests/fixtures/forms/peepl_biodata.docx"
    doc = Document()
    
    doc.add_heading("Biodata Form", 0)
    
    doc.add_paragraph("Nama Lengkap: ")
    doc.add_paragraph("NIK: ")
    doc.add_paragraph("Alamat: ")
    
    # Unicode symbol checkbox
    doc.add_paragraph("Jenis Kelamin: ☐ Laki-laki   ☐ Perempuan")
    doc.add_paragraph("Status Perkawinan: ☐ Kawin   ☐ Belum Kawin")
    
    # Table with labels and empty cells or checkboxes
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # Set labels and fields in cells
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Pekerjaan"
    hdr_cells[1].text = ""  # target empty field
    
    row_cells = table.rows[1].cells
    row_cells[0].text = "Agama"
    row_cells[1].text = ""  # target empty field
    
    row_cells2 = table.rows[2].cells
    row_cells2[0].text = "WNI"
    row_cells2[1].text = "☐ Ya   ☐ Tidak"
    
    doc.save(docx_path)
    print(f"Generated DOCX fixture: {docx_path}")

def generate_acroform_pdf():
    pdf_path = "tests/fixtures/forms/acroform_sample.pdf"
    width, height = letter
    c = canvas.Canvas(pdf_path, pagesize=letter)
    
    c.drawString(100, height - 100, "Interactive AcroForm Sample")
    c.drawString(100, height - 150, "Full Name:")
    c.drawString(100, height - 200, "Email:")
    
    form = c.acroForm
    form.textfield(
        name="full_name",
        x=200,
        y=height - 155,
        width=300,
        height=20,
        textColor=colors.black,
        borderColor=colors.black
    )
    form.textfield(
        name="email",
        x=200,
        y=height - 205,
        width=300,
        height=20,
        textColor=colors.black,
        borderColor=colors.black
    )
    
    c.save()
    print(f"Generated AcroForm PDF fixture: {pdf_path}")

def generate_static_pdf():
    pdf_path = "tests/fixtures/forms/static_sample.pdf"
    width, height = letter
    c = canvas.Canvas(pdf_path, pagesize=letter)
    
    c.drawString(100, height - 100, "Static PDF Form Sample (No Interactive Fields)")
    c.drawString(100, height - 150, "Name:")
    c.drawString(100, height - 200, "Address:")
    
    # Lines representing blanks for overlay filling
    c.setStrokeColor(colors.gray)
    # Lines representing blanks for overlay filling
    c.line(200, height - 150, 500, height - 150)
    c.line(200, height - 200, 500, height - 200)
    
    c.save()
    print(f"Generated Static PDF fixture: {pdf_path}")

def generate_xlsx():
    import openpyxl
    xlsx_path = "tests/fixtures/forms/sample_form.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Form Sheet"
    
    # Horizontal field alignment
    ws.cell(row=2, column=2).value = "Nama Lengkap:"
    # ws.cell(row=2, column=3).value is None (empty input field!)
    
    # Vertical field alignment
    ws.cell(row=4, column=2).value = "Alamat:"
    # ws.cell(row=5, column=2).value is None (empty input field!)
    
    wb.save(xlsx_path)
    print(f"Generated XLSX fixture: {xlsx_path}")

if __name__ == "__main__":
    setup_directories()
    generate_docx()
    generate_acroform_pdf()
    generate_static_pdf()
    generate_xlsx()
