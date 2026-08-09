import os
from docx import Document
from src.analyzers.docx_analyzer import DocxAnalyzer
from src.writers.docx_writer import DocxWriter

def test_docx_layout_preservation():
    doc_path = "tests/fixtures/forms/peepl_biodata.docx"
    output_path = "tests/fixtures/forms/peepl_biodata_layout_test.docx"
    
    original_doc = Document(doc_path)
    original_tables_count = len(original_doc.tables)
    original_paragraphs_count = len(original_doc.paragraphs)
    original_table_dims = [(len(t.rows), len(t.columns)) for t in original_doc.tables]
    
    analyzer = DocxAnalyzer()
    writer = DocxWriter()
    
    fields = analyzer.analyze(doc_path)
    
    # Fill all fields with dummy text/checks to verify formatting/layout preservation
    for f in fields:
        f.answer = "Test Value"
        
    writer.fill(doc_path, fields, output_path)
    
    filled_doc = Document(output_path)
    filled_tables_count = len(filled_doc.tables)
    filled_paragraphs_count = len(filled_doc.paragraphs)
    filled_table_dims = [(len(t.rows), len(t.columns)) for t in filled_doc.tables]
    
    if os.path.exists(output_path):
        os.remove(output_path)
        
    assert original_tables_count == filled_tables_count, "Number of tables changed."
    assert original_paragraphs_count == filled_paragraphs_count, "Number of paragraphs changed."
    assert original_table_dims == filled_table_dims, "Table dimensions changed."
