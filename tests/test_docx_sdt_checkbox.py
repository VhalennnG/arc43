import os
import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from src.analyzers.docx_analyzer import DocxAnalyzer
from src.writers.docx_writer import DocxWriter
from src.fields.models import FieldType, ResolutionMethod

def create_sdt_checkbox_document(doc_path: str):
    doc = Document()
    p = doc.add_paragraph("Status Pernikahan: ")
    
    # Construct structured document tag (w:sdt)
    sdt = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')
    
    # Set alias (w:alias) for field labelling
    alias = OxmlElement('w:alias')
    alias.set(qn('w:val'), 'Menikah')
    sdtPr.append(alias)
    
    # Set tag (w:tag)
    tag = OxmlElement('w:tag')
    tag.set(qn('w:val'), 'marital_status_sdt')
    sdtPr.append(tag)
    
    # Add checkbox element (w14:checkbox) under word 2010 namespace
    # w14 namespace is resolved in word docs as http://schemas.microsoft.com/office/word/2010/wordml
    checkbox = OxmlElement('w14:checkbox')
    checked = OxmlElement('w14:checked')
    checked.set(qn('w14:val'), '0')
    checkbox.append(checked)
    sdtPr.append(checkbox)
    sdt.append(sdtPr)
    
    # Add content block (w:sdtContent) showing the symbol
    sdtContent = OxmlElement('w:sdtContent')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = "☐"
    r.append(t)
    sdtContent.append(r)
    sdt.append(sdtContent)
    
    p._p.append(sdt)
    doc.save(doc_path)

def test_sdt_checkbox_detection_and_mutator():
    doc_path = "tests/fixtures/forms/sdt_test_form.docx"
    output_path = "tests/fixtures/forms/sdt_test_form_filled.docx"
    
    os.makedirs(os.path.dirname(doc_path), exist_ok=True)
    create_sdt_checkbox_document(doc_path)
    
    analyzer = DocxAnalyzer()
    writer = DocxWriter()
    
    # 1. Analyze
    fields = analyzer.analyze(doc_path)
    print("DETECTED FIELDS IN TEST:", fields)
    assert len(fields) == 1
    
    field = fields[0]
    assert field.field_type == FieldType.CHECKBOX
    assert field.checkbox_kind == "sdt"
    assert field.label == "Menikah"
    
    # Provide checked answer
    field.answer = "yes"
    
    # 2. Fill (mutate XML)
    writer.fill(doc_path, fields, output_path)
    
    # Load filled document and inspect raw XML elements
    filled_doc = Document(output_path)
    sdt_elements = filled_doc.element.xpath('.//*[local-name()="sdt"]')
    assert len(sdt_elements) == 1
    
    sdt = sdt_elements[0]
    
    # Verify XML checked state value is changed to '1'
    checked_nodes = sdt.xpath('.//*[local-name()="checked"]')
    assert len(checked_nodes) == 1
    
    checked_node = checked_nodes[0]
    val_attr = None
    for key in checked_node.attrib.keys():
        if key.endswith("val"):
            val_attr = key
            break
    assert val_attr is not None
    assert checked_node.get(val_attr) == "1"
    
    # Verify display character was swapped to checked character
    t_nodes = sdt.xpath('.//*[local-name()="sdtContent"]//*[local-name()="t"]')
    assert len(t_nodes) == 1
    assert t_nodes[0].text == "☒"
    
    # Clean up files
    if os.path.exists(doc_path):
        os.remove(doc_path)
    if os.path.exists(output_path):
        os.remove(output_path)
